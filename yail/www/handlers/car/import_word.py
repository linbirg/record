import io
import os
from typing import List, Dict, Tuple
from docx import Document
from lib.logger import LOG_INFO, LOG_WARNING, LOG_DEBUG


def parse_word_table(file_bytes: bytes, file_name: str = "") -> Tuple[List[Dict], List[Dict]]:
    """
    解析 Word 文档中的表格，提取姓名和车牌号

    Args:
        file_bytes: Word 文件字节数据
        file_name: 文件名（用于日志）

    Returns:
        (valid_rows, error_rows)
        valid_rows: [{'name': str, 'carNo': str}, ...]
        error_rows: [{'row': int, 'message': str}, ...]
    """
    LOG_INFO(f"[parse_word_table] 开始解析文件: {file_name}, 大小: {len(file_bytes)} bytes")

    # 检查文件扩展名
    ext = os.path.splitext(file_name)[1].lower() if file_name else ""
    if ext == ".doc":
        raise ValueError("不支持 .doc 格式，请将文档另存为 .docx 格式后重试")
    elif ext not in (".docx", ""):
        raise ValueError(f"不支持的文件格式: {ext}，仅支持 .docx")

    doc = Document(io.BytesIO(file_bytes))

    if not doc.tables:
        LOG_WARNING("[parse_word_table] 文档中没有表格")
        raise ValueError("文档中没有找到表格")

    table = doc.tables[0]
    LOG_INFO(f"[parse_word_table] 表格行数: {len(table.rows)}")

    if len(table.rows) < 2:
        LOG_WARNING("[parse_word_table] 表格行数少于2")
        raise ValueError("文档中没有找到车辆数据")

    header_row = table.rows[0]
    header_texts = [cell.text.strip() for cell in header_row.cells]
    LOG_INFO(f"[parse_word_table] 表头内容: {header_texts}")

    name_col = -1
    carNo_col = -1

    for i, cell in enumerate(header_row.cells):
        text = cell.text.strip()
        if '姓名' in text:
            name_col = i
            LOG_INFO(f"[parse_word_table] 找到姓名列: index={i}, text={text}")
        elif '车牌号' in text:
            carNo_col = i
            LOG_INFO(f"[parse_word_table] 找到车牌号列: index={i}, text={text}")

    if name_col == -1 or carNo_col == -1:
        LOG_WARNING(f"[parse_word_table] 未找到姓名或车牌号列: name_col={name_col}, carNo_col={carNo_col}")
        raise ValueError("表格中未找到姓名和车牌号列")

    valid_rows = []
    error_rows = []

    for row_idx, row in enumerate(table.rows[1:], start=2):
        try:
            name = row.cells[name_col].text.strip()
            car_no = row.cells[carNo_col].text.strip()
            LOG_DEBUG(f"[parse_word_table] 第{row_idx}行: name={name}, carNo={car_no}")

            if not name and not car_no:
                continue

            if not name or not car_no:
                error_rows.append({
                    'row': row_idx,
                    'message': f"姓名或车牌号为空 (姓名:{name}, 车牌号:{car_no})"
                })
                continue

            valid_rows.append({'name': name, 'carNo': car_no})
        except Exception as e:
            error_rows.append({'row': row_idx, 'message': str(e)})
            LOG_WARNING(f"[parse_word_table] 第{row_idx}行解析异常: {str(e)}")

    LOG_INFO(f"[parse_word_table] 解析完成: valid={len(valid_rows)}, error={len(error_rows)}")
    return valid_rows, error_rows
